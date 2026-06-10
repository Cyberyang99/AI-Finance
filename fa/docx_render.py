"""极简 markdown → Word(.docx) 渲染器（python-docx）.

支持 report/vet 这类 LLM 输出的 md 子集：
  # ~ ###### 标题 / 段落 / - * 无序列表 / 1. 有序列表 / | 表格 |
  **粗体** *斜体* `代码` 行内样式 / > 引用 / --- 分隔线 / ``` 代码块
不是通用 md 解析器；解析不了的行按普通段落兜底，绝不丢内容。
"""

import re
from pathlib import Path

_INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*[^*\n]+?\*|`[^`\n]+?`)")


def _add_runs(paragraph, text: str) -> None:
    """把 **粗体**/*斜体*/`代码` 拆成 runs 写进段落。"""
    for token in _INLINE_RE.split(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**") and len(token) > 4:
            paragraph.add_run(token[2:-2]).bold = True
        elif token.startswith("`") and token.endswith("`") and len(token) > 2:
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
        elif token.startswith("*") and token.endswith("*") and len(token) > 2:
            paragraph.add_run(token[1:-1]).italic = True
        else:
            paragraph.add_run(token)


def _strip_inline(text: str) -> str:
    return re.sub(r"\*\*(.+?)\*\*|\*([^*\n]+?)\*|`([^`\n]+?)`",
                  lambda m: m.group(1) or m.group(2) or m.group(3), text)


def _is_table_sep(line: str) -> bool:
    body = line.strip().strip("|")
    return bool(body) and set(body) <= set("-: |")


def _disp_width(s: str) -> int:
    """显示宽度估算：CJK 算 2、其余算 1（列宽按内容比例分配用）。"""
    return sum(2 if ord(ch) > 0x2E7F else 1 for ch in s)


def _shade_cell(cell, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def md_to_docx(md: str, path: Path, title: str = "", meta: str = "") -> None:
    """渲染 markdown 到 .docx。title 为文档主标题，meta 为标题下的灰色说明行。"""
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    doc = Document()
    # Word 的拉丁字体不带中文映射，必须显式设 eastAsia 否则中文回退宋体
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    if title:
        doc.add_heading(title, level=0)
    if meta:
        p = doc.add_paragraph()
        run = p.add_run(meta)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    lines = md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # 代码块
        if stripped.startswith("```"):
            block = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            i += 1  # 收尾的 ```
            p = doc.add_paragraph()
            run = p.add_run("\n".join(block))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            continue

        # 表格块（首行 | 开头，次行是分隔行）
        if stripped.startswith("|") and i + 1 < len(lines) and _is_table_sep(lines[i + 1]):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                if not _is_table_sep(lines[i]):
                    cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                    rows.append(cells)
                i += 1
            if rows:
                n_cols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=n_cols)
                table.style = "Table Grid"
                # 列宽按内容显示宽度成比例分配（含表头），加上下限：
                # 短列（✓/评分）不塌缩、长文本列（结论/风险）拿大头，提升可读性
                weights = []
                for ci in range(n_cols):
                    lens = [_disp_width(r[ci]) for r in rows if ci < len(r) and r[ci]]
                    weights.append(min(max(max(lens) if lens else 4, 4), 40))
                sect = doc.sections[0]
                usable = int(sect.page_width - sect.left_margin - sect.right_margin)
                widths = [int(usable * w / sum(weights)) for w in weights]
                table.autofit = False
                for ci in range(n_cols):
                    table.columns[ci].width = widths[ci]
                for ri, row in enumerate(rows):
                    for ci in range(n_cols):
                        cell = table.cell(ri, ci)
                        cell.width = widths[ci]
                        text = row[ci] if ci < len(row) else ""
                        para = cell.paragraphs[0]
                        if ri == 0:
                            para.add_run(_strip_inline(text)).bold = True
                            _shade_cell(cell, "EFEFEF")
                        else:
                            _add_runs(para, text)
            continue

        # 标题
        m = re.match(r"^(#{1,6})\s+(.*)", stripped)
        if m:
            level = min(len(m.group(1)), 4)
            doc.add_heading(_strip_inline(m.group(2)).strip(), level=level)
            i += 1
            continue

        # 分隔线
        if re.fullmatch(r"(-{3,}|\*{3,}|_{3,})", stripped):
            i += 1
            continue

        # 引用
        if stripped.startswith(">"):
            p = doc.add_paragraph(style="Intense Quote")
            _add_runs(p, stripped.lstrip("> ").strip())
            i += 1
            continue

        # 列表
        m = re.match(r"^\s*[-*+]\s+(.*)", line)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, m.group(1))
            i += 1
            continue
        m = re.match(r"^\s*\d+[.、)]\s+(.*)", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            _add_runs(p, m.group(1))
            i += 1
            continue

        # 普通段落（兜底）
        p = doc.add_paragraph()
        _add_runs(p, stripped)
        i += 1

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
