"""Word 抽文 — python-docx + 直接读 XML 抓文本框.

研报常见坑：
1. 大量正文放在浮动文本框 (w:txbxContent) 里，python-docx 默认遍历不到
2. 页眉页脚里有标题/作者/免责声明，也要抓
3. 表格里嵌套表格（python-docx 自动处理，但格式要保住）
4. 备注/批注（w:comment）有时含分析师附注
"""

from pathlib import Path
from xml.etree import ElementTree as ET

W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _extract_textbox_texts(doc) -> list[str]:
    """python-docx 不暴露 textbox，必须直接读 document.xml。

    遍历 body 的所有 w:txbxContent 节点，把里面的 w:t 拼起来。
    """
    out = []
    try:
        body_xml = doc.element.body
    except Exception:
        return out

    for txbx in body_xml.iter(f"{W_NS}txbxContent"):
        # 收集每个段落 (w:p) 内的所有 w:t
        for p in txbx.iter(f"{W_NS}p"):
            text_parts = [t.text for t in p.iter(f"{W_NS}t") if t.text]
            line = "".join(text_parts).strip()
            if line:
                out.append(line)
    return out


def _extract_header_footer(doc) -> list[str]:
    """页眉页脚（每节可不同）。"""
    out = []
    for section in doc.sections:
        for region in (section.header, section.footer,
                       section.first_page_header, section.first_page_footer,
                       section.even_page_header, section.even_page_footer):
            try:
                for p in region.paragraphs:
                    t = p.text.strip()
                    if t:
                        out.append(t)
            except Exception:
                continue
    return out


def load_docx(path: Path) -> tuple[str, int]:
    """返回 (纯文本, 段落数)。表格按 markdown 表格输出。

    抽取范围:
      - 正文段落（含标题分级）
      - 正文表格（markdown）
      - 浮动文本框（直接读 XML）
      - 页眉页脚（去重）
    """
    from docx import Document

    doc = Document(str(path))
    parts = []
    para_count = 0

    # 1) 正文段落
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        # 程序生成的 docx 段落可能引用未定义样式 → p.style 为 None，
        # 用 getattr 安全取名，取不到就当普通正文处理（不丢内容）。
        style = (getattr(p.style, "name", None) or "").lower()
        if "heading 1" in style:
            parts.append(f"# {text}")
        elif "heading 2" in style:
            parts.append(f"## {text}")
        elif "heading" in style:
            parts.append(f"### {text}")
        else:
            parts.append(text)
        para_count += 1

    # 2) 正文表格
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            rows.append("| " + " | ".join(cells) + " |")
        if rows:
            sep = "| " + " | ".join(["---"] * len(table.rows[0].cells)) + " |"
            rows.insert(1, sep)
            parts.append("\n".join(rows))
            para_count += len(rows)

    # 3) 浮动文本框（很多研报核心论点在这）
    tb_lines = _extract_textbox_texts(doc)
    if tb_lines:
        seen = set()
        unique_tb = []
        for line in tb_lines:
            if line not in seen:
                seen.add(line)
                unique_tb.append(line)
        parts.append("\n\n## [文本框内容]\n\n" + "\n\n".join(unique_tb))
        para_count += len(unique_tb)

    # 4) 页眉页脚（去重 + 跳过同文）
    hf_lines = _extract_header_footer(doc)
    if hf_lines:
        seen = set()
        unique_hf = []
        for line in hf_lines:
            if line not in seen and len(line) > 4:
                seen.add(line)
                unique_hf.append(line)
        if unique_hf:
            parts.append("\n\n## [页眉页脚]\n\n" + "\n".join(unique_hf))

    return "\n\n".join(parts), para_count
