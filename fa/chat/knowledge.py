"""Knowledge query helpers for fa chat.

This module keeps the interactive entrypoint deterministic: search CoT + notes,
render full matched content, and optionally export a Word document to Desktop.
It does not mutate memory.
"""

from __future__ import annotations

import re
from difflib import SequenceMatcher
from datetime import datetime
from pathlib import Path
from typing import Any, Optional

from ..cot import load_cots
from ..ingest.user_note import load_user_notes
from ..memory.store import PROJECT_DIR
from ..sectors import resolve_alias, resolve_theme_tag
from .resolver import _normalize_ticker, resolve


def _desktop_dir() -> Path:
    return Path.home() / "Desktop"


def _norm_text(text: str) -> str:
    return re.sub(r"\s+", "", (text or "").lower())


def _query_terms(query: str) -> list[str]:
    """Split natural-language search into useful terms.

    Mixed strings like "token聚合平台" become ["token", "聚合平台", "聚合", "平台"].
    This keeps exact keyword search useful while avoiding the brittle "whole
    phrase must appear" behavior.
    """
    raw = re.findall(r"[a-zA-Z0-9]+|[\u4e00-\u9fff]+", (query or "").lower())
    out: list[str] = []

    def add(term: str):
        term = term.strip().lower()
        if term and term not in out:
            out.append(term)

    for term in raw:
        add(term)
        # Chinese chunks often arrive unsegmented. Add stable 2-char phrases and
        # edge 3-char phrases without introducing too much generic noise.
        if re.fullmatch(r"[\u4e00-\u9fff]+", term) and len(term) >= 3:
            for i in range(len(term) - 1):
                add(term[i:i + 2])
            if len(term) >= 4:
                add(term[:3])
                add(term[-3:])
    return out


def _score_text(query: str, terms: list[str], hay: str) -> float:
    """Loose relevance score for natural-language memory search."""
    if not terms:
        return 1.0
    hay_l = (hay or "").lower()
    hay_n = _norm_text(hay)
    query_n = _norm_text(query)
    score = 0.0
    weak_terms = {"ai", "平台", "公司", "行业", "业务", "逻辑", "相关", "投资", "市场"}

    if query_n and query_n in hay_n:
        score += 12.0

    for term in terms:
        term_n = _norm_text(term)
        if not term_n:
            continue
        if term_n in hay_n:
            # Longer exact matches carry more information. Two-char Chinese grams
            # still count, but less than full phrases.
            if term_n in weak_terms:
                score += 0.8
            else:
                score += 2.0 + min(len(term_n), 8) * 0.25
            continue

        # Soft fallback for near terms. Keep this light to avoid turning every
        # generic Chinese character overlap into a hit.
        if len(term_n) >= 4:
            ratio = SequenceMatcher(None, term_n, hay_n[:4000]).quick_ratio()
            if ratio >= 0.18:
                score += ratio

    return score


def _min_relevance(terms: list[str]) -> float:
    if not terms:
        return 0.0
    # Multi-term natural queries should not match on one generic word.
    return 3.8 if len(terms) >= 3 else 2.0


def _resolve_query_ticker(raw: str) -> Optional[str]:
    if not raw:
        return None
    raw = raw.strip()
    if re.match(r"^\d{1,6}\.(HK|SHE|SHG|SZ|SH|US)$", raw, re.I):
        return _normalize_ticker(raw)
    hits = resolve(raw, limit=1)
    return hits[0]["ticker"] if hits else None


def _clean_note_body(content: str) -> str:
    return content.split("---", 2)[-1].strip() if content else ""


def _cot_hay(c: dict) -> str:
    return " ".join([
        str(c.get("trigger", "")),
        str(c.get("COT", "")),
        str(c.get("evidence", "")),
        str(c.get("_source", "")),
        str(c.get("_sector", "")),
        str(c.get("_ticker", "")),
        " ".join(c.get("_tags") or []),
    ])


def _note_hay(n: dict) -> str:
    return " ".join([
        str(n.get("ticker", "")),
        str(n.get("created_at", "")),
        str(n.get("content", "")),
        str(n.get("path", "")),
    ])


def query_knowledge(
    query: str = "",
    scope: str = "all",
    ticker: str = "",
    tag: str = "",
    sector: str = "",
    min_signal: int = 0,
    max_items: int = 20,
    full: bool = True,
    export_docx: bool = False,
) -> dict[str, Any]:
    """Search CoT and/or note memory.

    Returns {hits, total, text, docx_path, warnings}. Hits are full records
    suitable for rendering/export. This is read-only except optional docx export.
    """
    scope = (scope or "all").lower()
    if scope not in {"all", "cot", "note"}:
        scope = "all"
    max_items = max(1, min(int(max_items or 20), 200))
    min_signal = max(0, int(min_signal or 0))

    warnings: list[str] = []
    terms = _query_terms(query)
    resolved_ticker = _resolve_query_ticker(ticker) if ticker else None

    resolved_tag = ""
    invalid_filter = False
    if tag:
        resolved_tag, cands = resolve_theme_tag(tag)
        if not resolved_tag:
            warnings.append(f"主题 '{tag}' 未匹配；候选: {'、'.join(cands[:12])}")
            invalid_filter = True

    resolved_sector = ""
    if sector:
        resolved_sector = resolve_alias(sector) or sector

    hits: list[dict[str, Any]] = []

    if invalid_filter:
        text = render_knowledge_hits([], total=0, query=query or tag, full=full, warnings=warnings)
        return {"hits": [], "total": 0, "text": text, "docx_path": "", "warnings": warnings}

    if scope in {"all", "cot"}:
        cots = load_cots(sector=resolved_sector or None,
                         min_signal=min_signal,
                         tag=resolved_tag or None)
        for c in cots:
            hay = _cot_hay(c)
            if resolved_ticker:
                bound = c.get("_ticker") == resolved_ticker
                mentioned = resolved_ticker.lower() in hay.lower()
                if not (bound or mentioned):
                    continue
            relevance = _score_text(query, terms, hay)
            if terms and relevance < _min_relevance(terms):
                continue
            try:
                signal = int(c.get("signal", 0) or 0)
            except (TypeError, ValueError):
                signal = 0
            hits.append({
                "kind": "cot",
                "score": signal,
                "rank_score": relevance * 100 + signal,
                "title": c.get("trigger", ""),
                "id": c.get("_cot_id", ""),
                "ticker": c.get("_ticker", ""),
                "date": c.get("_created_at", ""),
                "source": c.get("_source", ""),
                "sector": c.get("_sector", ""),
                "tags": c.get("_tags") or [],
                "signal": c.get("signal", ""),
                "body": c.get("COT", "").strip(),
                "evidence": c.get("evidence", "").strip(),
            })

    if scope in {"all", "note"}:
        notes = load_user_notes(ticker=resolved_ticker) if resolved_ticker else load_user_notes()
        for n in notes:
            relevance = _score_text(query, terms, _note_hay(n))
            if terms and relevance < _min_relevance(terms):
                continue
            body = _clean_note_body(n.get("content", ""))
            title = f"{n.get('ticker', '')} note {n.get('created_at', '')}"
            hits.append({
                "kind": "note",
                "score": 0,
                "rank_score": relevance * 100,
                "title": title,
                "id": str(n.get("path", "")),
                "ticker": n.get("ticker", ""),
                "date": n.get("created_at", ""),
                "source": Path(str(n.get("path", ""))).name,
                "sector": "",
                "tags": [],
                "signal": "",
                "body": body,
                "evidence": "",
            })

    hits.sort(key=lambda h: (-(h.get("rank_score") or 0), 0 if h["kind"] == "cot" else 1, h.get("date") or ""))
    total = len(hits)
    selected = hits[:max_items]
    text = render_knowledge_hits(selected, total=total, query=query, full=full, warnings=warnings)
    docx_path = export_hits_docx(selected, query=query, total=total, warnings=warnings) if export_docx and selected else None
    return {
        "hits": selected,
        "total": total,
        "text": text,
        "docx_path": str(docx_path) if docx_path else "",
        "warnings": warnings,
    }


def render_knowledge_hits(
    hits: list[dict[str, Any]],
    total: Optional[int] = None,
    query: str = "",
    full: bool = True,
    warnings: Optional[list[str]] = None,
) -> str:
    warnings = warnings or []
    total = len(hits) if total is None else total
    title = f"=== 知识库查询：{query or '全库'}（命中 {total}，展示 {len(hits)}）==="
    lines = [title]
    for w in warnings:
        lines.append(f"⚠ {w}")
    if not hits:
        lines.append("没有命中。可换关键词，或改用主题/ticker 查询。")
        return "\n".join(lines)

    for i, h in enumerate(hits, 1):
        if h["kind"] == "cot":
            tags = "、".join(h.get("tags") or []) or "(未打主题)"
            lines.append(
                f"\n[{i}] CoT [{h.get('signal', '?')}/10] {h.get('title', '')}\n"
                f"    id={h.get('id', '')} | 主题={tags} | 行业={h.get('sector', '')} | 来源={h.get('source', '')}"
            )
            if full:
                lines.append(f"    推理链: {h.get('body', '')}")
                if h.get("evidence"):
                    lines.append(f"    原文依据: {h['evidence']}")
        else:
            lines.append(
                f"\n[{i}] note {h.get('ticker', '')} {h.get('date', '')}\n"
                f"    来源={h.get('source', '')}"
            )
            if full:
                lines.append(h.get("body", ""))

    if total > len(hits):
        lines.append(f"\n... 还有 {total - len(hits)} 条未展示。可缩小关键词或导出 Word 查看。")
    return "\n".join(lines)


def export_hits_docx(
    hits: list[dict[str, Any]],
    query: str = "",
    total: Optional[int] = None,
    warnings: Optional[list[str]] = None,
    output_path: Optional[Path] = None,
) -> Path:
    from docx import Document
    from docx.shared import Pt

    warnings = warnings or []
    total = len(hits) if total is None else total
    stamp = datetime.now().strftime("%Y%m%d_%H%M")
    slug = re.sub(r"[\\/:*?\"<>|\s]+", "_", (query or "all")[:24]).strip("_") or "all"
    output_path = output_path or (_desktop_dir() / f"fa_knowledge_query_{slug}_{stamp}.docx")

    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)

    doc.add_heading("fa 知识库查询", level=0)
    doc.add_paragraph(f"查询：{query or '全库'}")
    doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(f"命中总数：{total}；本文收录：{len(hits)}")
    for w in warnings:
        doc.add_paragraph(f"警告：{w}")

    for i, h in enumerate(hits, 1):
        if h["kind"] == "cot":
            doc.add_heading(f"{i}. CoT [{h.get('signal', '?')}/10] {h.get('title', '')}", level=1)
            tags = "、".join(h.get("tags") or []) or "(未打主题)"
            doc.add_paragraph(
                f"id={h.get('id', '')} | 主题={tags} | 行业={h.get('sector', '')} | 来源={h.get('source', '')}"
            )
            doc.add_paragraph("推理链：")
            doc.add_paragraph(h.get("body", ""))
            if h.get("evidence"):
                doc.add_paragraph(f"原文依据：{h['evidence']}")
        else:
            doc.add_heading(f"{i}. note {h.get('ticker', '')} {h.get('date', '')}", level=1)
            doc.add_paragraph(f"来源={h.get('source', '')}")
            for block in re.split(r"\n{2,}", h.get("body", "")):
                if block.strip():
                    doc.add_paragraph(block.strip())

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
